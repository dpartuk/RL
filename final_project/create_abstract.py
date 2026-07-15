"""Generate the project abstract as a formatted .docx following the template."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURE_PATH = os.path.join(OUTPUT_DIR, "experiment_1d", "output_plots", "4_trajectories.png")
OUT_PATH = os.path.join(OUTPUT_DIR, "abstract.docx")

doc = Document()

# --- Page margins: 2.5 cm all round ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Helper to set font on a run ---
def set_run_font(run, name, size, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font for East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def set_paragraph_spacing(para, before=0, after=0, line_spacing=1.0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

# --- Title (Arial 14pt, bold, centred) ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Dropout-Aware Q-Learning for Trajectory Tracking Under "
    "Lossy Network Conditions: A Simplified Simulation Study Inspired "
    "by Inverse Reinforcement Learning With Data Dropouts"
)
set_run_font(run, "Arial", 14, bold=True)
set_paragraph_spacing(title, before=0, after=6)

# --- Author (Times New Roman 14pt) ---
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = author.add_run("*Peleg, D.")
set_run_font(run, "Times New Roman", 14)
set_paragraph_spacing(author, before=0, after=0)

# --- Lead presenter ---
presenter = doc.add_paragraph()
presenter.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = presenter.add_run("*lead presenter")
set_run_font(run, "Times New Roman", 10)
set_paragraph_spacing(presenter, before=0, after=0)

# --- Affiliation ---
affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = affil.add_run("doron.peleg@gmail.com, Afeka College of Engineering, Tel Aviv, Israel")
set_run_font(run, "Times New Roman", 10)
set_paragraph_spacing(affil, before=0, after=12)

# --- Body paragraphs (Times New Roman 12pt, justified, single-spaced) ---
body_paragraphs = [
    (
        "Standard reinforcement learning (RL) algorithms assume that the agent "
        "receives reliable state observations at every time step. In real-world "
        "networked control systems, however, sensor readings and control signals "
        "are transmitted over wireless channels where packets are randomly lost. "
        "Fan et al. (2025) addressed this challenge in their paper \"Inverse "
        "Reinforcement Learning for Discrete-Time Systems With Data Dropouts\" "
        "(IEEE Transactions on Cybernetics, Vol. 55, No. 4), proposing both a "
        "model-based IRL algorithm using a Smith predictor for state estimation "
        "and a data-driven, model-free inverse Q-learning algorithm that "
        "explicitly accounts for random state dropouts. Their work provides "
        "rigorous convergence proofs and demonstrates that learning both the "
        "cost function and the control policy from incomplete data outperforms "
        "methods that rely on manually specified cost functions."
    ),
    (
        "This project presents a simplified, educational simulation that "
        "captures the core insights of Fan et al. We construct a 1D trajectory "
        "tracking environment in which an agent must follow a sinusoidal target "
        "using three discrete actions (left, stay, right), with 30% of state "
        "observations randomly dropped to simulate lossy network conditions. "
        "Three tabular Q-learning agents are compared: (1) a Baseline agent "
        "trained with perfect observations, (2) a Naive agent that trains under "
        "dropout but treats stale (last-received) states as current, and "
        "(3) a Dropout-Aware agent that skips Q-table updates whenever the "
        "current state is known to be stale. Through systematic ablation "
        "testing, we found that the dominant source of performance degradation "
        "is Q-table corruption from writing updates to incorrect state entries, "
        "and that simply withholding updates during dropout episodes is the "
        "most effective countermeasure—reducing the training reward gap relative "
        "to the naive agent by approximately 19% (from −45.9 to −43.2 vs. the "
        "baseline’s −31.8)."
    ),
    (
        "Additionally, we implement an inverse RL pipeline that recovers the "
        "reward function solely from expert demonstrations, without access to "
        "the hand-coded reward. Using a five-dimensional feature representation "
        "(bullseye indicator, linear penalty, quadratic penalty, danger zone, "
        "and bias), the IRL gradient-descent loop converges within 200 "
        "iterations and produces a policy whose evaluation performance is "
        "identical to the hand-coded baseline (mean tracking error = 0.086, "
        "total reward = −25.7), confirming that the expert’s implicit objective "
        "was successfully recovered. A full 3×2 comparison (three agent "
        "variants × two reward sources) demonstrates that the IRL-learned "
        "reward generalises across all agent types, preserving the same "
        "performance ordering observed under the hand-coded reward."
    ),
    (
        "Our findings reinforce the paper’s central message: explicitly "
        "accounting for data dropouts in the learning process yields measurable "
        "improvements, and inverse RL provides a viable alternative to manual "
        "reward engineering. The full implementation (1,139 lines of Python, "
        "NumPy, and Matplotlib) and all experimental results are publicly "
        "available at github.com/dpartuk/RL."
    ),
]

for text in body_paragraphs:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run_font(run, "Times New Roman", 12)
    set_paragraph_spacing(para, before=0, after=6)

# --- Figure (if exists) ---
if os.path.exists(FIGURE_PATH):
    # Add figure
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(FIGURE_PATH, width=Cm(12))
    set_paragraph_spacing(fig_para, before=6, after=2)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(
        "Figure 1. Agent trajectories vs. sinusoidal target under 30% "
        "packet dropout. Top: Baseline (no dropout). Middle: Naive and "
        "Dropout-Aware agents. Bottom: Dropout mask (red = packet lost)."
    )
    set_run_font(run, "Times New Roman", 10, italic=True)
    set_paragraph_spacing(cap, before=0, after=6)

# --- Reference ---
ref_para = doc.add_paragraph()
ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = ref_para.add_run(
    "Reference: Fan, J., Shi, P., Xue, W., Lian, B., Cui, Y., & Lewis, "
    "F. L. (2025). Inverse Reinforcement Learning for Discrete-Time "
    "Systems With Data Dropouts. IEEE Transactions on Cybernetics, 55(4), "
    "1744–1757."
)
set_run_font(run, "Times New Roman", 10)
set_paragraph_spacing(ref_para, before=0, after=0)

doc.save(OUT_PATH)
print(f"Abstract saved to: {OUT_PATH}")
